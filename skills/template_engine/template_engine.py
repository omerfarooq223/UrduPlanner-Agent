"""
UrduPlanner — Word template parser and filler.

Understands the exact lesson plan (سبق کی منصوبہ بندی) table layout:
  - Each table = 1 lesson (14 rows × 4 cols)
  - Col 3 = labels (static), Cols 0-2 = merged content (dynamic)
  - 3 tables per week (فی ہفتہ سبق کی تعداد: ۳)

Row mapping:
  0  = Header (سبق کی منصوبہ بندی / session year)
  1  = Lessons/week, Subject, Teaching week, Dates
  2  = Duration, Class, Unit number, Dates
  3  = Lesson title + keywords (سبق کا عنوان)
  4  = Learning outcomes (حاصلات تعلم)
  5  = Teaching resources (تدریسی وسائل)
  6  = Teaching method (طریقہ تدریس)
  7  = Introduction / warm-up questions (تعارف / آغاز)
  8  = Core teaching activities (بنیادی تدریس)
  9  = Classwork / reading exercises (کلاس ورک)
  10 = Closing activity / questions (اختتامی سرگرمی)
  11 = Assessment (تشخیصی عمل)
  12 = Homework (گھر کا کام) — may be blank
  13 = Teaching review (جائزہ تدریس) — may be blank
"""

import copy
import os

from typing import Union
from docx import Document
from docx.document import Document as DocxDocument
from docx.oxml.ns import qn


URDU_FONT_NAME = "Urdu Typesetting"


# Row indices in each lesson table
ROW_HEADER     = 0
ROW_META1      = 1   # lessons/week, subject, teaching week, dates
ROW_META2      = 2   # duration, class, unit number, dates
ROW_TITLE      = 3   # lesson title + keywords
ROW_OUTCOMES   = 4   # learning outcomes
ROW_RESOURCES  = 5   # teaching resources
ROW_METHOD     = 6   # teaching method label row
ROW_INTRO      = 7   # introduction / warm-up questions
ROW_CORE       = 8   # core teaching activities
ROW_CLASSWORK  = 9   # classwork / reading exercises
ROW_CLOSING    = 10  # closing discussion questions
ROW_ASSESSMENT = 11  # assessment
ROW_HOMEWORK   = 12  # homework (optional)
ROW_REVIEW     = 13  # teaching review (optional)

# Fields that the LLM generates per lesson
DYNAMIC_FIELDS = [
    "teaching_week",   # row 1, col 2 — e.g. تدریسی ہفتہ: ۸
    "dates",           # row 1, col 3 — e.g. تاریخ:: ۹مارچ تا ۱۳مارچ
    "unit_number",     # row 2, col 2 — e.g. یونٹ نمبر: ۴
    "title",           # row 3, cols 0-2 — lesson title + keywords
    "outcomes",        # row 4, cols 0-2 — learning outcomes
    "resources",       # row 5, cols 0-2 — teaching resources
    "intro",           # row 7, cols 0-2 — intro questions + duration
    "core_teaching",   # row 8, cols 0-2 — core teaching activities + duration
    "classwork",       # row 9, cols 0-2 — reading/writing exercises + duration
    "closing",         # row 10, cols 0-2 — closing questions + duration
    "assessment",      # row 11, cols 0-2 — assessment description
    "homework",        # row 12, cols 0-2 — homework (can be empty)
    "review",          # row 13, cols 0-2 — teaching review (can be empty)
]


def read_template(template_path: str) -> Document:
    """Load the .docx template."""
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Template not found: {template_path}")
    return Document(template_path)


def get_template_structure(doc: Document) -> dict:
    """
    Extract the template structure showing what each lesson table looks like.
    Returns a dict with row labels and content from table 0.

    NOTE: This is a debugging/inspection utility only.
    Do NOT pass its output to the LLM — style guidance is hardcoded
    in SYSTEM_PROMPT in content_generator.py.
    """
    structure = {
        "num_lessons": len(doc.tables),
        "rows_per_lesson": 14,
        "sample_lesson": {},
    }

    if doc.tables:
        table = doc.tables[0]
        for r, row in enumerate(table.rows):
            label = row.cells[3].text.strip() if len(row.cells) > 3 else ""
            content = row.cells[0].text.strip()
            structure["sample_lesson"][f"row_{r}"] = {
                "label": label,
                "content": content,
            }

    return structure


def fill_lesson_table(table, lesson_data: dict):
    """
    Fill one lesson table with generated content.

    lesson_data keys match DYNAMIC_FIELDS:
      teaching_week, dates, unit_number, title, outcomes, resources,
      intro, core_teaching, classwork, closing, assessment, homework, review
    """
    # Bug fix: removed "dates2" — it was never in lesson_data so never triggered.
    # The manual copy below handles the duplicate dates in ROW_META2.
    field_to_row = {
        "teaching_week": (ROW_META1,    [2]),
        "dates":         (ROW_META1,    [3]),
        "unit_number":   (ROW_META2,    [2]),
        "title":         (ROW_TITLE,    [0, 1, 2]),
        "outcomes":      (ROW_OUTCOMES, [0, 1, 2]),
        "resources":     (ROW_RESOURCES,[0, 1, 2]),
        "intro":         (ROW_INTRO,    [0, 1, 2]),
        "core_teaching": (ROW_CORE,     [0, 1, 2]),
        "classwork":     (ROW_CLASSWORK,[0, 1, 2]),
        "closing":       (ROW_CLOSING,  [0, 1, 2]),
        "assessment":    (ROW_ASSESSMENT,[0, 1, 2]),
        "homework":      (ROW_HOMEWORK, [0, 1, 2]),
        "review":        (ROW_REVIEW,   [0, 1, 2]),
    }

    for field, value in lesson_data.items():
        if field not in field_to_row:
            continue
        row_idx, cols = field_to_row[field]
        for col in cols:
            cell = table.rows[row_idx].cells[col]
            _replace_cell_text(cell, value)

    # Copy dates into ROW_META2 col 3 — template has date in both meta rows
    if "dates" in lesson_data:
        _replace_cell_text(
            table.rows[ROW_META2].cells[3],
            lesson_data["dates"],
        )


def fill_all_lessons(
    template_path: Union[str, DocxDocument],
    lessons: list[dict],
    output_path: str,
) -> str:
    """
    Fill all lesson tables in the template.

    template_path: either a file path string or an already-loaded Document object.
    lessons: list of dicts, one per table/lesson, with keys from DYNAMIC_FIELDS.
    """
    # Bug fix: accept both a path string and a pre-loaded Document object
    # so web flow (app.py) and CLI flow (main.py) can both call this correctly.
    doc = template_path if isinstance(template_path, DocxDocument) else Document(template_path)

    if len(lessons) > len(doc.tables):
        raise ValueError(
            f"Template has {len(doc.tables)} tables but got {len(lessons)} lessons"
        )

    for i, lesson_data in enumerate(lessons):
        fill_lesson_table(doc.tables[i], lesson_data)

    os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
    doc.save(output_path)
    return output_path


def _replace_paragraph_text(para, new_text: str):
    """Replace text in a paragraph while keeping the first run's formatting."""
    if not para.runs:
        para.text = new_text
        return

    first_run = para.runs[0]
    for run in para.runs:
        run.text = ""
    first_run.text = new_text


def _replace_cell_text(cell, new_text: str):
    """
    Replace all text in a table cell while preserving RTL formatting.

    Bug fix: clear the cell contents first so template paragraphs do not leak
    through into the generated document, then recreate only the paragraphs
    needed for the new content.
    """
    template_para = cell.paragraphs[0] if cell.paragraphs else None
    template_run = template_para.runs[0] if template_para and template_para.runs else None
    template_pPr = copy.deepcopy(template_para._element.pPr) if template_para and template_para._element.pPr is not None else None

    # Remove the existing cell content but keep the cell properties / borders.
    if hasattr(cell._tc, "clear_content"):
        cell._tc.clear_content()
    else:
        for child in list(cell._tc):
            if child.tag.endswith("}tcPr"):
                continue
            cell._tc.remove(child)

    lines = new_text.split("\n") if new_text else [""]

    # Write only the new paragraphs that the field actually needs.
    for line in lines:
        new_para = cell.add_paragraph()

        if template_para is not None and template_para.style is not None:
            new_para.style = template_para.style

        # Copy paragraph properties: RTL direction, spacing, alignment.
        if template_pPr is not None:
            new_para._element.insert(0, copy.deepcopy(template_pPr))

        # Add run and copy run properties: font, size, bold, color.
        new_run = new_para.add_run(line)
        if template_run is not None:
            new_run.bold = template_run.bold
            new_run.italic = template_run.italic
            new_run.underline = template_run.underline
            new_run.font.name = template_run.font.name
            new_run.font.size = template_run.font.size
            new_run.font.rtl = template_run.font.rtl
            if template_run.font.color is not None:
                new_run.font.color.rgb = template_run.font.color.rgb
        _force_urdu_font(new_run)


def _force_urdu_font(run):
    """Force the run font to Urdu Typesetting for Urdu script consistency."""
    run.font.name = URDU_FONT_NAME
    run.font.rtl = True
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    r_fonts.set(qn("w:ascii"), URDU_FONT_NAME)
    r_fonts.set(qn("w:hAnsi"), URDU_FONT_NAME)
    r_fonts.set(qn("w:cs"), URDU_FONT_NAME)